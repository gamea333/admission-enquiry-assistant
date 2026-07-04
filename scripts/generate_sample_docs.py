"""Generate reproducible sample documents for the admission enquiry assistant."""

from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCUMENTS_DIR = PROJECT_ROOT / "data" / "documents"

SCHOOL_NAME = "Greenfield International School"
ACADEMIC_YEAR = "2026-27"


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=14,
            textColor=colors.HexColor("#1a5276"),
        ),
        "heading": ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor("#2874a6"),
        ),
        "body": ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=8,
        ),
    }


def _build_pdf(filename: str, story: list) -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCUMENTS_DIR / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.build(story)
    return path


def generate_prospectus() -> Path:
    s = _pdf_styles()
    story = [
        Paragraph(f"{SCHOOL_NAME} — School Prospectus {ACADEMIC_YEAR}", s["title"]),
        Paragraph(
            f"Established in 1998, {SCHOOL_NAME} is a CBSE-affiliated co-educational institution "
            "serving Nursery through Grade 10 in Bengaluru. We nurture curious, compassionate learners "
            "through inquiry-based teaching and strong values.",
            s["body"],
        ),
        Paragraph("Our Mission", s["heading"]),
        Paragraph(
            "To provide holistic education that develops academic excellence, critical thinking, "
            "and responsible citizenship in a safe, inclusive environment.",
            s["body"],
        ),
        Paragraph("Campus & Facilities", s["heading"]),
        Paragraph(
            "• 8-acre green campus with CCTV surveillance and visitor management<br/>"
            "• Smart classrooms with interactive panels (Grades 3–10)<br/>"
            "• Science, robotics, and computer laboratories<br/>"
            "• Library with 18,000+ volumes and digital reading zone<br/>"
            "• Semi-Olympic swimming pool, basketball courts, and indoor sports hall<br/>"
            "• School infirmary with qualified nurse on campus (8 AM – 3 PM)<br/>"
            "• GPS-enabled air-conditioned transport fleet (35 routes)",
            s["body"],
        ),
        Paragraph("Curriculum Overview", s["heading"]),
        Paragraph(
            "<b>Early Years (Nursery – UKG):</b> Play-way methodology, phonics, numeracy, art, and motor skills.<br/>"
            "<b>Primary (Grades 1–5):</b> CBSE foundational literacy and numeracy; EVS, languages, ICT basics.<br/>"
            "<b>Middle (Grades 6–8):</b> Subject-specialist teaching; introduction to third language from Grade 6.<br/>"
            "<b>Senior (Grades 9–10):</b> Board-focused preparation with remedial support and career counselling.",
            s["body"],
        ),
        Paragraph("Co-curricular & Student Life", s["heading"]),
        Paragraph(
            "Students participate in house competitions, Model UN, coding club, music, dance, "
            "and inter-school sports. Annual events include Science Expo, Literary Fest, and Sports Day.",
            s["body"],
        ),
        Paragraph("Contact", s["heading"]),
        Paragraph(
            "Admissions Office: admissions@greenfield.edu.in | +91 80 4123 5600<br/>"
            "Campus: 42 Lakeview Road, Whitefield, Bengaluru – 560066",
            s["body"],
        ),
    ]
    return _build_pdf("prospectus.pdf", story)


def generate_admission_policy() -> Path:
    s = _pdf_styles()
    story = [
        Paragraph(f"{SCHOOL_NAME} — Admission Policy {ACADEMIC_YEAR}", s["title"]),
        Paragraph("Eligibility Criteria", s["heading"]),
        Paragraph(
            "<b>Nursery:</b> Child must be 3 years old as on 1 June 2026.<br/>"
            "<b>LKG:</b> 4 years as on 1 June 2026.<br/>"
            "<b>UKG:</b> 5 years as on 1 June 2026.<br/>"
            "<b>Grade 1:</b> 6 years as on 1 June 2026; basic reading readiness assessed.<br/>"
            "<b>Grades 2–9:</b> Previous class report card; transfer certificate required for Grades 2+.<br/>"
            "<b>Grade 10:</b> Admission only if seats available; prior CBSE/ICSE enrolment preferred.",
            s["body"],
        ),
        Paragraph("Admission Process", s["heading"]),
        Paragraph(
            "1. <b>Online Registration</b> — Complete the form at admissions.greenfield.edu.in and pay "
            "₹1,500 registration fee (non-refundable).<br/>"
            "2. <b>Document Upload</b> — Birth certificate, Aadhaar, passport-size photo, previous report card.<br/>"
            "3. <b>Interaction / Assessment</b> — Nursery–UKG: parent-child interaction; Grades 1–10: written assessment.<br/>"
            "4. <b>Provisional Offer</b> — Email within 5 working days of assessment.<br/>"
            "5. <b>Fee Payment & Confirmation</b> — Pay admission fee and deposit within 7 days to confirm seat.<br/>"
            "6. <b>Orientation</b> — Mandatory parent orientation in the last week of May 2026.",
            s["body"],
        ),
        Paragraph("Required Documents", s["heading"]),
        Paragraph(
            "• Original birth certificate and one photocopy<br/>"
            "• Aadhaar card of child and both parents<br/>"
            "• Previous 2 years' report cards (if applicable)<br/>"
            "• Transfer certificate (Grades 2 and above)<br/>"
            "• Migration certificate (if joining from another board)<br/>"
            "• Medical fitness certificate<br/>"
            "• 4 passport-size photographs",
            s["body"],
        ),
        Paragraph("Important Dates", s["heading"]),
        Paragraph(
            "<b>Registration opens:</b> 1 December 2025<br/>"
            "<b>Early-bird registration closes:</b> 31 January 2026<br/>"
            "<b>Assessment window:</b> 15 February – 30 April 2026<br/>"
            "<b>General registration closes:</b> 15 May 2026 (subject to seat availability)<br/>"
            "<b>Session begins:</b> 2 June 2026",
            s["body"],
        ),
        Paragraph("General Policies", s["heading"]),
        Paragraph(
            "Admissions are granted based on age eligibility, assessment performance, and seat availability. "
            "The school reserves the right to refuse admission without assigning reasons. "
            "Sibling applicants receive priority during tie-break situations. "
            "Fees once paid are governed by the school's refund policy (see fee structure document).",
            s["body"],
        ),
    ]
    return _build_pdf("admission_policy.pdf", story)


def generate_fee_structure() -> Path:
    s = _pdf_styles()
    fee_rows = [
        ["Grade", "Admission Fee (₹)", "Tuition Fee / Year (₹)", "Transport Fee / Year (₹)", "One-time Deposit (₹)"],
        ["Nursery", "25,000", "1,20,000", "36,000", "15,000"],
        ["LKG", "25,000", "1,25,000", "36,000", "15,000"],
        ["UKG", "25,000", "1,30,000", "36,000", "15,000"],
        ["Grade 1", "30,000", "1,45,000", "38,000", "20,000"],
        ["Grade 2", "30,000", "1,50,000", "38,000", "20,000"],
        ["Grade 3", "30,000", "1,55,000", "38,000", "20,000"],
        ["Grade 4", "30,000", "1,60,000", "38,000", "20,000"],
        ["Grade 5", "35,000", "1,70,000", "40,000", "20,000"],
        ["Grade 6", "35,000", "1,80,000", "40,000", "25,000"],
        ["Grade 7", "35,000", "1,85,000", "40,000", "25,000"],
        ["Grade 8", "40,000", "1,95,000", "42,000", "25,000"],
        ["Grade 9", "40,000", "2,05,000", "42,000", "25,000"],
        ["Grade 10", "45,000", "2,15,000", "42,000", "30,000"],
    ]

    table = Table(fee_rows, colWidths=[1.0 * inch, 1.15 * inch, 1.35 * inch, 1.35 * inch, 1.35 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2874a6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#eaf2f8")]),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    story = [
        Paragraph(f"{SCHOOL_NAME} — Fee Structure {ACADEMIC_YEAR}", s["title"]),
        Paragraph(
            "All fees are payable annually unless stated otherwise. Transport fee applies only if availing school bus. "
            "One-time deposit is refundable when the student leaves, subject to clearance of dues.",
            s["body"],
        ),
        table,
        Spacer(1, 12),
        Paragraph("Payment Notes", s["heading"]),
        Paragraph(
            "• Tuition fee may be paid in two instalments (June and November) with 2% surcharge on second instalment.<br/>"
            "• Sibling discount: 10% on tuition fee for the second child, 15% for the third.<br/>"
            "• Late payment attracts ₹500 per week after the due date.<br/>"
            "• Uniform, books, and meals are charged separately.",
            s["body"],
        ),
    ]
    return _build_pdf("fee_structure.pdf", story)


def generate_faq_docx() -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCUMENTS_DIR / "faq.docx"

    faqs = [
        (
            "Uniforms",
            "What is the school uniform policy?",
            "Students must wear the prescribed uniform on all school days: grey trousers/skirt, "
            "white shirt, school tie, black shoes, and school belt. House T-shirts are worn on Wednesdays. "
            "Sports uniform is required on days with PE periods. Uniforms are available at the school store "
            "and authorised vendor on MG Road.",
        ),
        (
            "Holidays",
            "How many holidays does the school have in a year?",
            "The academic calendar includes approximately 110 instructional days per term with breaks for "
            "Dasara (1 week), Winter (2 weeks), and Summer (6 weeks). A detailed holiday list is shared "
            "in the parent handbook at the start of each session.",
        ),
        (
            "Fees",
            "Is there a sibling discount on fees?",
            "Yes. A 10% tuition fee discount applies to the second sibling and 15% to the third sibling "
            "enrolled concurrently. The discount does not apply to admission fee, transport, or deposit.",
        ),
        (
            "Extracurriculars",
            "What extracurricular activities are offered?",
            "The school offers cricket, football, swimming, chess, robotics, classical dance, western music, "
            "debate, and art club. Activities run after school from 3:15–4:30 PM on designated weekdays. "
            "Registration opens in June each year.",
        ),
        (
            "Campus Visit",
            "How can I book a campus visit?",
            "Campus tours are available on Saturdays between 9 AM and 12 PM. Book online at "
            "admissions.greenfield.edu.in/visit or call +91 80 4123 5601. Walk-ins are accommodated "
            "subject to slot availability. Tours last approximately 45 minutes.",
        ),
        (
            "Transport",
            "Can I change my child's bus route mid-year?",
            "Route changes are permitted at the start of each term (June and November) if seats are available "
            "on the requested route. Submit a written request to the transport office at least 15 days in advance.",
        ),
        (
            "Meals",
            "Does the school provide lunch?",
            "Optional hot lunch is available through our catering partner on a monthly subscription basis. "
            "Vegetarian and Jain meal options are offered. Students may also bring home-packed lunch.",
        ),
        (
            "Assessment",
            "Is there an entrance test for Nursery admission?",
            "No written test for Nursery–UKG. A brief parent-child interaction assesses readiness and "
            "communication. Grades 1 and above require a written assessment in English and Mathematics.",
        ),
        (
            "Communication",
            "How does the school communicate with parents?",
            "We use the Greenfield Parent App for daily updates, circulars, and fee reminders. "
            "Parent-teacher meetings are held every term. Urgent alerts are sent via SMS.",
        ),
        (
            "Withdrawal",
            "What is the withdrawal and refund policy?",
            "One month's written notice is required for withdrawal. Tuition fee is refunded on a pro-rata "
            "basis for unused months. Admission fee is non-refundable. Deposit is refunded after TC issuance "
            "and clearance of all dues.",
        ),
    ]

    doc = Document()
    doc.add_heading(f"{SCHOOL_NAME} — Parent FAQs", level=0)
    doc.add_paragraph(f"Academic Year {ACADEMIC_YEAR}")

    for category, question, answer in faqs:
        doc.add_heading(question, level=2)
        doc.add_paragraph(f"Category: {category}")
        doc.add_paragraph(answer)

    doc.save(path)
    return path


def generate_school_timings() -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCUMENTS_DIR / "school_timings.txt"

    content = f"""{SCHOOL_NAME} — Class-wise School Timings ({ACADEMIC_YEAR})
================================================================

EARLY YEARS SECTION (Nursery, LKG, UKG)
-----------------------------------------
School Hours     : 8:30 AM – 12:30 PM (Monday to Friday)
Gate Opens       : 8:15 AM
Dispersal Window : 12:20 PM – 12:45 PM
Saturday         : No regular classes (optional activity camps announced separately)

PRIMARY SECTION (Grades 1 – 5)
-------------------------------
School Hours     : 8:00 AM – 2:30 PM (Monday to Friday)
Assembly         : 8:00 AM – 8:15 AM
Lunch Break      : 12:30 PM – 1:00 PM
Dispersal        : 2:15 PM – 2:45 PM
Saturday         : 8:00 AM – 12:00 PM (Grades 4–5: remedial / club activities, as scheduled)

MIDDLE SECTION (Grades 6 – 8)
------------------------------
School Hours     : 7:45 AM – 3:00 PM (Monday to Friday)
Assembly         : 7:45 AM – 8:00 AM
Lunch Break      : 12:45 PM – 1:15 PM
Dispersal        : 2:45 PM – 3:15 PM
Saturday         : 7:45 AM – 12:30 PM (sports / lab sessions, rotating schedule)

SENIOR SECTION (Grades 9 – 10)
-------------------------------
School Hours     : 7:30 AM – 3:15 PM (Monday to Friday)
Assembly         : 7:30 AM – 7:45 AM
Lunch Break      : 12:30 PM – 1:00 PM
Dispersal        : 3:00 PM – 3:30 PM
Saturday         : 7:30 AM – 1:00 PM (board exam preparation sessions)

GENERAL NOTES
-------------
• All students must be collected within 30 minutes of dispersal time.
• Late pickup after 3:45 PM (Primary/Middle/Senior) incurs ₹100 per 15-minute block.
• Early dismissal requests must be submitted via the Parent App before 10:00 AM.
• School remains closed on all gazetted public holidays per Karnataka state calendar.

Transport Pickup: Buses depart campus 15 minutes after section dispersal time.
Contact Transport Desk: transport@greenfield.edu.in | +91 80 4123 5602
"""
    path.write_text(content, encoding="utf-8")
    return path


def main() -> None:
    generated = [
        generate_prospectus(),
        generate_admission_policy(),
        generate_fee_structure(),
        generate_faq_docx(),
        generate_school_timings(),
    ]
    print(f"Generated {len(generated)} files in {DOCUMENTS_DIR}:")
    for p in generated:
        print(f"  - {p.name}")


if __name__ == "__main__":
    main()
