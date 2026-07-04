import shutil
import sys
from collections import Counter
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _normalize_metadata(doc: Document, source_file: str, source_type: str) -> Document:
    page_number = doc.metadata.get("page")
    if page_number is None:
        page_number = doc.metadata.get("page_number")

    doc.metadata = {
        "source_file": source_file,
        "source_type": source_type,
        "page_number": page_number,
    }
    return doc


def _load_pdf(file_path: Path) -> list[Document]:
    docs = PyPDFLoader(str(file_path)).load()
    return [_normalize_metadata(doc, file_path.name, "pdf") for doc in docs]


def _load_docx(file_path: Path) -> list[Document]:
    docx = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
    if not paragraphs:
        return []

    return [
        Document(
            page_content="\n".join(paragraphs),
            metadata={
                "source_file": file_path.name,
                "source_type": "docx",
                "page_number": None,
            },
        )
    ]


def _load_txt(file_path: Path) -> list[Document]:
    docs = TextLoader(str(file_path), encoding="utf-8").load()
    return [_normalize_metadata(doc, file_path.name, "txt") for doc in docs]


def load_documents(documents_dir: Path | None = None) -> list[Document]:
    settings = get_settings()
    docs_path = documents_dir or settings.documents_dir
    docs_path.mkdir(parents=True, exist_ok=True)

    documents: list[Document] = []
    loaders = {
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".txt": _load_txt,
    }

    for file_path in sorted(docs_path.iterdir()):
        if not file_path.is_file():
            continue
        loader = loaders.get(file_path.suffix.lower())
        if loader is None:
            continue
        documents.extend(loader(file_path))

    return documents


def chunk_documents(documents: list[Document]) -> list[Document]:
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    return splitter.split_documents(documents)


def get_embeddings() -> GoogleGenerativeAIEmbeddings:
    settings = get_settings()
    return GoogleGenerativeAIEmbeddings(
        model=settings.embedding_model,
        google_api_key=settings.google_api_key,
    )


def _reset_vectorstore(vectorstore_dir: Path) -> None:
    if vectorstore_dir.exists():
        shutil.rmtree(vectorstore_dir)
    vectorstore_dir.mkdir(parents=True, exist_ok=True)


def build_vectorstore(documents_dir: Path | None = None) -> tuple[Chroma, Counter[str]]:
    settings = get_settings()
    documents = load_documents(documents_dir)
    if not documents:
        raise ValueError(f"No supported documents found in {settings.documents_dir}")

    chunks = chunk_documents(documents)
    chunk_counts = Counter(chunk.metadata["source_file"] for chunk in chunks)

    _reset_vectorstore(settings.vectorstore_dir)
    embeddings = get_embeddings()

    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(settings.vectorstore_dir),
    )
    return vectorstore, chunk_counts


def get_vectorstore() -> Chroma:
    settings = get_settings()
    settings.vectorstore_dir.mkdir(parents=True, exist_ok=True)
    return Chroma(
        persist_directory=str(settings.vectorstore_dir),
        embedding_function=get_embeddings(),
    )


def ingest(documents_dir: Path | None = None) -> Counter[str]:
    settings = get_settings()
    if not settings.google_api_key:
        raise ValueError("GOOGLE_API_KEY is not set. Add it to your .env file.")

    _, chunk_counts = build_vectorstore(documents_dir)
    return chunk_counts


def _print_summary(chunk_counts: Counter[str], total_chunks: int) -> None:
    print(f"\nIngestion complete — {total_chunks} chunks embedded.\n")
    print("Chunks per source file:")
    for source_file, count in sorted(chunk_counts.items()):
        print(f"  {source_file}: {count}")


if __name__ == "__main__":
    try:
        counts = ingest()
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)

    _print_summary(counts, sum(counts.values()))
